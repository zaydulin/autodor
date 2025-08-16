import io
import json
import os
from datetime import datetime
from decimal import Decimal

from django.conf import settings
from django.core.exceptions import PermissionDenied
from django.core.files.base import ContentFile
from django.utils import timezone
from django.utils.decorators import method_decorator
from django.views import View
from docx import Document
from moderation.tasks import start_call_task, end_call_task

from django.contrib.auth.decorators import login_required
from django.db import models, transaction
from django.http import JsonResponse, HttpResponse, HttpResponseServerError
from django.shortcuts import render, get_object_or_404
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_POST
from django.views.generic import ListView, DetailView, TemplateView, FormView
from django.db.models import Q
from django.contrib.auth.mixins import LoginRequiredMixin
from .models import AdvertAplication, ChatMessage, CallSession, AdvertDocument, AdvertExpense, BaseDocument
from moderation.models import Advert, AdvertAplication
from webmain.models import Faqs, Seo

from useraccount.models import Profile


class AdvertAplicationListView(LoginRequiredMixin, ListView):
    model = AdvertAplication
    template_name = "site/useraccount/advertaplication.html"
    context_object_name = "advertaplications"
    paginate_by = 20

    def get_queryset(self):
        # фильтруем M2M по текущему пользователю
        return (
            AdvertAplication.objects.filter(user=self.request.user)
            .select_related("advert")
            .prefetch_related("user")
            .order_by("-created_at")
        )

class AdvertAplicationDetailView(LoginRequiredMixin, DetailView):
    model = AdvertAplication
    template_name = "site/useraccount/advertaplication-detail.html"
    context_object_name = "application"

    def get_queryset(self):
        return (
            super().get_queryset()
            .filter(user=self.request.user)
            .select_related("advert")
            .prefetch_related("user")
        )

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        application = self.object
        advert = application.advert

        # добавляем объявление
        context["application"] = application
        context["advert"] = advert
        expenses = application.expenses.all()
        context['expenses'] = expenses
        total_expenses = sum(expense.amount for expense in expenses)
        users_list = (
                list(application.user.all()) +
                list(application.user_menager.all()) +
                list(application.user_drivers.all())
        )
        context['users'] = [user for user in users_list if user != self.request.user]
        context['total_price'] = application.price + advert.price + total_expenses

        user = application.user.first()
        messages = ChatMessage.objects.filter(
            applications=application
        ).filter(
            Q(author=user) |
            Q(author__in=application.user_menager.all()) |
            Q(author__in=application.user_drivers.all())
        ).order_by('date')  # сортируем по времени
        context['messages'] = messages

        calls = CallSession.objects.filter(application=application)
        context['documents'] = application.documents.all().order_by('-created_at')
        context['calls'] = calls

        context['all_managers'] = Profile.objects.filter(type=0,employee=2)
        context['all_drivers'] = Profile.objects.filter(type=0,employee=1)

        return context


@csrf_exempt
def update_application(request, application_id):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            application = AdvertAplication.objects.get(id=application_id)
            updated_fields = {}

            # Получаем текущих пользователей
            current_users = set(application.user.all())

            # Обработка полей
            for key, value in data.items():
                if hasattr(application, key):
                    field_obj = getattr(application, key)

                    # Для обычных полей
                    if isinstance(field_obj, (str, int, float, type(None))):
                        setattr(application, key, value)
                        updated_fields[key] = str(getattr(application, key))

                    # Для DecimalField
                    elif isinstance(field_obj, models.DecimalField):
                        setattr(application, key, value)
                        updated_fields[key] = str(getattr(application, key))

                    # Для ManyToMany полей
                    elif isinstance(getattr(application, key), models.fields.related_descriptors.ManyToManyDescriptor):
                        if value is None:
                            continue

                        if isinstance(value, list):
                            related_manager = getattr(application, key)

                            if len(value) > 0:
                                related_manager.set(value)

                                # Добавляем новых пользователей в общий список
                                if key == 'user_menager' or key == 'user_drivers':
                                    new_users = User.objects.filter(id__in=value)
                                    current_users.update(new_users)
                            else:
                                related_manager.clear()

                            updated_fields[key] = ','.join(str(v) for v in value)

            # Обновляем общий список пользователей
            application.user.set(current_users)
            application.save()

            return JsonResponse({'success': True, 'updated_fields': updated_fields})

        except AdvertAplication.DoesNotExist:
            return JsonResponse({'success': False, 'error': 'Заявка не найдена'}, status=404)
        except Exception as e:
            return JsonResponse({'success': False, 'error': str(e)}, status=500)

    return JsonResponse({'success': False})


# Create your views here.

def _to_int(v):
    try:
        return int(v)
    except (TypeError, ValueError):
        return None

def _to_decimal(v):
    try:
        # точка/запятая не важны
        return float(str(v).replace(',', '.'))
    except (TypeError, ValueError):
        return None

class AdvertView(ListView):
    template_name = 'site/useraccount/adverts.html'
    context_object_name = 'adverts'
    model = Advert
    paginate_by = 16

    def get_queryset(self):
        qs = Advert.objects.all().order_by('-created_at')
        g = self.request.GET

        # Поиск
        q = g.get('q')
        if q:
            qs = qs.filter(
                Q(name__icontains=q) |
                Q(subtitle__icontains=q) |
                Q(article__icontains=q) |
                Q(description__icontains=q) |
                Q(brand__icontains=q) |
                Q(model_auto__icontains=q) |
                Q(color__icontains=q)
            )

        # Марка/модель
        brand = g.get('brand')
        if brand:
            qs = qs.filter(brand__iexact=brand)

        model_auto = g.get('model_auto')
        if model_auto:
            qs = qs.filter(model_auto__iexact=model_auto)

        # Валюта
        currency = g.get('currency')
        if currency:
            qs = qs.filter(currency__iexact=currency)

        # Цена
        price_min = _to_decimal(g.get('price_min'))
        price_max = _to_decimal(g.get('price_max'))
        if price_min is not None:
            qs = qs.filter(price__gte=price_min)
        if price_max is not None:
            qs = qs.filter(price__lte=price_max)

        # Год
        year_min = _to_int(g.get('year_min'))
        year_max = _to_int(g.get('year_max'))
        if year_min is not None:
            qs = qs.filter(year__gte=year_min)
        if year_max is not None:
            qs = qs.filter(year__lte=year_max)

        # Пробег
        mileage_min = _to_int(g.get('mileage_min'))
        mileage_max = _to_int(g.get('mileage_max'))
        if mileage_min is not None:
            qs = qs.filter(mileage__gte=mileage_min)
        if mileage_max is not None:
            qs = qs.filter(mileage__lte=mileage_max)

        # Мощность
        power_min = _to_int(g.get('power_min'))
        power_max = _to_int(g.get('power_max'))
        if power_min is not None:
            qs = qs.filter(power__gte=power_min)
        if power_max is not None:
            qs = qs.filter(power__lte=power_max)

        # Объем двигателя
        ev_min = _to_decimal(g.get('engine_volume_min'))
        ev_max = _to_decimal(g.get('engine_volume_max'))
        if ev_min is not None:
            qs = qs.filter(engine_volume__gte=ev_min)
        if ev_max is not None:
            qs = qs.filter(engine_volume__lte=ev_max)

        # Двери
        doors = _to_int(g.get('doors'))
        if doors is not None:
            qs = qs.filter(doors=doors)

        # Цвет
        color = g.get('color')
        if color:
            qs = qs.filter(color__icontains=color)

        # Коробка/топливо/привод (множественный выбор)
        transmissions = g.getlist('transmission')
        if transmissions:
            qs = qs.filter(transmission__in=transmissions)

        fuels = g.getlist('fuel')
        if fuels:
            qs = qs.filter(fuel__in=fuels)

        drives = g.getlist('drive')
        if drives:
            qs = qs.filter(drive__in=drives)

        # Есть изображения
        has_images = g.get('has_images')
        if has_images == '1':
            qs = qs.exclude(images__isnull=True).exclude(images=[])

        # Сортировка
        order = g.get('order')
        if order == 'price_asc':
            qs = qs.order_by('price', '-created_at')
        elif order == 'price_desc':
            qs = qs.order_by('-price', '-created_at')
        elif order == 'year_desc':
            qs = qs.order_by('-year', '-created_at')
        elif order == 'year_asc':
            qs = qs.order_by('year', '-created_at')
        elif order == 'mileage_asc':
            qs = qs.order_by('mileage', '-created_at')
        elif order == 'mileage_desc':
            qs = qs.order_by('-mileage', '-created_at')
        else:
            # по умолчанию — свежие
            qs = qs.order_by('-created_at')

        return qs

    def get_context_data(self, **kwargs):
        ctx = super().get_context_data(**kwargs)
        g = self.request.GET

        ctx['brands'] = (Advert.objects.values_list('brand', flat=True)
                         .exclude(brand__isnull=True).exclude(brand__exact='')
                         .distinct().order_by('brand'))
        ctx['models'] = (Advert.objects.values_list('model_auto', flat=True)
                         .exclude(model_auto__isnull=True).exclude(model_auto__exact='')
                         .distinct().order_by('model_auto'))
        ctx['currencies'] = (Advert.objects.values_list('currency', flat=True)
                             .exclude(currency__isnull=True).exclude(currency__exact='')
                             .distinct().order_by('currency'))
        ctx['transmission_choices'] = Advert.TransmissionType.choices
        ctx['fuel_choices'] = Advert.FuelType.choices
        ctx['drive_choices'] = Advert.DriveType.choices

        # чтобы в шаблоне не вызывать getlist(...)
        ctx['selected_transmissions'] = g.getlist('transmission')
        ctx['selected_fuels'] = g.getlist('fuel')
        ctx['selected_drives'] = g.getlist('drive')

        # для остальных полей оставим доступ к params.*
        ctx['params'] = g
        return ctx


class AdvertDetailView(DetailView):
    """Страница новости"""
    model = Advert
    template_name = 'site/useraccount/adverts_detail.html'
    context_object_name = 'advert'
    slug_field = "pk"


"""ЧаВо"""
class FaqsModerView(ListView):
    model = Faqs
    template_name = 'site/useraccount/faqs.html'  # No .html extension
    context_object_name = 'faqs'
    paginate_by = 10

    def get_queryset(self):
        return Faqs.objects.filter(publishet=True)

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        try:
            seo_data = Seo.objects.get(pagetype=4)
            context['seo_previev'] = seo_data.previev
            context['seo_title'] = seo_data.title
            context['seo_description'] = seo_data.metadescription
            context['seo_propertytitle'] = seo_data.propertytitle
            context['seo_propertydescription'] = seo_data.propertydescription
        except Seo.DoesNotExist:
            context['seo_previev'] = None
            context['seo_title'] = None
            context['seo_description'] = None
            context['seo_propertytitle'] = None
            context['seo_propertydescription'] = None

        return context


def create_application(request, advert_id):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            advert = Advert.objects.get(id=advert_id)
            svoi_price = Decimal(data.get('price'))

            application = AdvertAplication.objects.create(
                advert=advert,
                status=AdvertAplication.Status.NEW,
                price = svoi_price,
            )
            docunments_base = BaseDocument.objects.all().count()
            for document in range(docunments_base):
                AdvertDocument.objects.create(
                    aplication=application,
                    document_type=2,
                )

            application.user.add(request.user)
            application.save()

            return JsonResponse({'success': True, 'application_id': str(application.id)})

        except Advert.DoesNotExist:
            return JsonResponse({'success': False, 'error': 'Объявление не найдено'})
        except Exception as e:
            return JsonResponse({'success': False, 'error': str(e)})
    return JsonResponse({'success': False, 'error': 'Invalid request method'})




@login_required
@require_POST
def send_message(request, app_id):
    application = get_object_or_404(AdvertAplication, id=app_id)
    content = request.POST.get('content')
    if content:
        message = ChatMessage.objects.create(
            applications=application,
            content=content,
            author=request.user,
        )
        return JsonResponse({
            'id': str(message.id),
            'author': message.author.username,
            'content': message.content,
            'date': message.date.strftime('%Y-%m-%d %H:%M:%S'),
        })
    return JsonResponse({'error': 'Нет содержимого'}, status=400)


@login_required
def get_new_messages(request, app_id):
    application = get_object_or_404(AdvertAplication, id=app_id)
    last_message_id = request.GET.get('last_message_id')
    current_user = request.user

    if last_message_id:
        try:
            last_message = ChatMessage.objects.get(id=last_message_id)
            new_messages = ChatMessage.objects.filter(
                applications=application,
                date__gt=last_message.date
            ).exclude(author=current_user).order_by('date')
        except ChatMessage.DoesNotExist:
            new_messages = ChatMessage.objects.filter(
                applications=application
            ).exclude(author=current_user).order_by('date')
    else:
        new_messages = ChatMessage.objects.filter(
            applications=application
        ).exclude(author=current_user).order_by('date')

    messages_data = []
    for msg in new_messages:
        messages_data.append({
            'id': str(msg.id),
            'author': msg.author.username,
            'content': msg.content,
            'date': msg.date.strftime('%Y-%m-%d %H:%M:%S'),
        })

    return JsonResponse({'messages': messages_data})


def start_call(request, application_id):
    if request.method == 'POST':
        try:
            application = get_object_or_404(AdvertAplication, id=application_id)
            call_session = CallSession.objects.create(
                application=application,
                caller=request.user,
                callee=application.user.first()
            )
            call_id = call_session.id
            start_call_task.delay(call_id)
            return JsonResponse({'call_id': call_id})
        except Exception as e:
            import logging
            logger = logging.getLogger(__name__)
            logger.exception("Error in start_call")
            return JsonResponse({'error': str(e)}, status=500)
    return JsonResponse({'error': 'Invalid request'}, status=400)


def generate_contract(request, application_id):
    # Получаем данные
    application = AdvertAplication.objects.get(id=application_id)
    advert = application.advert

    # Создаем документ с нуля
    doc = Document()
    doc.add_heading('ДОГОВОР КУПЛИ-ПРОДАЖИ ТОВАРА', level=0)

    # Дата
    doc.add_paragraph(f"Дата: {timezone.now().strftime('%d.%m.%Y')}")

    # Описание ТС
    doc.add_paragraph("Передаваемое транспортное средство:")
    table = doc.add_table(rows=0, cols=2)
    def add_row(label, value):
        row = table.add_row()
        row.cells[0].text = label
        row.cells[1].text = str(value) if value is not None else ''

    add_row('Марка', advert.brand or '')
    add_row('Модель', advert.model_auto or '')
    add_row('Год выпуска', advert.year or '')
    add_row('Пробег', advert.mileage or '')
    add_row('Цвет', advert.color or '')
    add_row('Объем двигателя', advert.engine_volume or '')
    add_row('Мощность', advert.power or '')
    add_row('Тип КПП', advert.get_transmission_display())
    add_row('Топливо', advert.get_fuel_display())
    add_row('Привод', advert.get_drive_display())
    add_row('Адрес размещения', advert.address or '')
    add_row('Цена', application.price)

    # Доп. информация может быть добавлена по желанию
    # Примерно — можно добавить подписи, условия и т.д.

    # Сохраняем в BytesIO
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    file_name = f"contract_{application.id}.docx"
    content = ContentFile(buffer.getvalue(), name=file_name)

    with transaction.atomic():
        advert_document = AdvertDocument.objects.create(
            aplication=application,
            document_type=AdvertDocument.DocumentType.CONTRACT,
            file=content
        )

    # Возвращаем файл пользователю (или можно вернуть путь/пометить как созданный)
    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="{file_name}"'
    return response



@login_required
def call_page(request, application_id,calle_id):
    application = get_object_or_404(AdvertAplication, id=application_id)
    calle = Profile.objects.get(id=calle_id)
    call, created = CallSession.objects.get_or_create(
        application=application,
        defaults={
            'caller': request.user,
            'callee': calle
        }
    )

    if not call.callee:
        return HttpResponse("Нет пользователя для звонка", status=400)

    other_user = call.callee if request.user == call.caller else call.caller


    return render(request, 'site/useraccount/call_page.html', {
        'application_id': application_id,
        'call_id': str(call.id),
        'user': request.user,
        'other_user': other_user,
    })


@method_decorator(login_required, name='dispatch')
class CreateExpenseView(View):
    def post(self, request):
        try:
            # Парсим JSON данные
            data = json.loads(request.body)

            # Получаем данные из запроса
            application_id = data.get('application')
            title = data.get('title')
            amount = data.get('amount')
            date = datetime.now()

            # Проверяем права доступа
            application = AdvertAplication.objects.get(id=application_id)
            if request.user not in application.user_menager.all():
                raise PermissionDenied("У вас нет прав на добавление расходов")

            # Создаем новый расход
            expense = AdvertExpense.objects.create(
                aplication=application,
                title=title,
                amount=amount,
                date=date
            )

            return JsonResponse({
                'success': True,
                'message': 'Расход успешно добавлен',
                'expense_id': expense.id
            }, status=201)


        except PermissionDenied as e:
            return JsonResponse({
                'success': False,
                'message': str(e)
            }, status=403)

